"""Build the Description of Improvements document from Mark's own template.

The template at app/templates/Improvements.docx is the layout and the truth.
This module reads its shape rather than carrying a list of headings in code, so
changing a label in Word changes the output with no code change. That is the
whole point of the template being a real file Spenser can open.

What the template looks like, and therefore what this reads:

    heading   a paragraph whose whole text is bold and ends with a colon
    field     a paragraph starting with a bold "Label:" then a tab, then the
              value. The template ships these with the value empty
    prose     an empty paragraph inside a block that has no fields

One block heading carries the placeholder [NAME]. That block repeats once per
tenancy, per building or per use, and the placeholder becomes the tenancy's own
name. Blaul Lofts uses three: Common Areas, Commercial Suite, Apartment Units.

A field with no value does not silently vanish. It is written with the words
that say so, because a blank costs Mark ten seconds and an invented value
reaches his client over his signature.
"""
import copy
import re
from pathlib import Path

from docx import Document

PLACEHOLDER = "[NAME]"
NOT_FOUND = "[not in the assessor card or the transcript]"

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _text(para):
    return "".join(node.text or "" for node in para._p.iter(_W + "t"))


def _leading_bold(para):
    """The bold run text at the start of a paragraph, or ''.

    Word splits a typed label across several runs, so this walks them while
    they stay bold instead of trusting the first one to hold the whole label.
    """
    out = ""
    for run in para.runs:
        if not run.bold:
            break
        out += run.text
    return out


def _has_tab(para) -> bool:
    """A tab inside a run, which is what separates a label from its value.

    Only runs. A tab stop declared in the paragraph's properties is the same
    XML tag and means something else entirely.
    """
    return any(run._r.find(_W + "tab") is not None for run in para.runs)


def is_heading(para) -> bool:
    label = _leading_bold(para).strip()
    if not label.endswith(":") or _has_tab(para):
        return False
    return label == _text(para).strip()


def is_field(para) -> bool:
    return _leading_bold(para).strip().endswith(":") and _has_tab(para)


def field_label(para) -> str:
    return _leading_bold(para).strip().rstrip(":").strip()


def read_shape(template_path: Path):
    """The template's own blocks, as [(heading, [field labels])]."""
    doc = Document(template_path)
    blocks = []
    for para in doc.paragraphs:
        if is_heading(para):
            blocks.append((_text(para).strip(), []))
        elif is_field(para) and blocks:
            blocks[-1][1].append(field_label(para))
    return blocks


def _clear_value(para):
    """Drop every run after the label and its tab."""
    keep = 0
    for run in para.runs:
        keep += 1
        if "\t" in run.text:
            break
    for run in list(para.runs)[keep:]:
        run._r.getparent().remove(run._r)


def _set_value(para, value: str):
    _clear_value(para)
    para.add_run(value)


def _block_paragraphs(doc, heading_text: str):
    """Every paragraph of one block, its heading first."""
    out, inside = [], False
    for para in doc.paragraphs:
        if is_heading(para):
            if inside:
                break
            inside = _text(para).strip() == heading_text
        if inside:
            out.append(para)
    return out


def _repeat_interior(doc, names):
    """Clone the [NAME] block once per name, in order, and drop the original.

    Cloning the paragraphs rather than building new ones keeps whatever styling
    the template actually carries, including anything Spenser changes in Word
    later that this module has never heard of.
    """
    heading = next((p for p in doc.paragraphs if PLACEHOLDER in _text(p)), None)
    if heading is None:
        return []
    block = _block_paragraphs(doc, _text(heading).strip())
    anchor = block[-1]._p
    made = []
    for name in names:
        copies = []
        for para in block:
            new = copy.deepcopy(para._p)
            anchor.addnext(new)
            anchor = new
            copies.append(new)
        # A blank line between one tenancy's block and the next.
        gap = copy.deepcopy(block[0]._p)
        for child in list(gap):
            if child.tag != _W + "pPr":
                gap.remove(child)
        anchor.addnext(gap)
        anchor = gap
        made.append((name, copies))
    for para in block:
        para._p.getparent().remove(para._p)
    return made


def _fill_placeholder(p_element, name: str):
    for node in p_element.iter(_W + "t"):
        if node.text and PLACEHOLDER in node.text:
            node.text = node.text.replace(PLACEHOLDER, name)


def build_improvements_docx(values: dict, template_path: Path, out_path: Path,
                            interiors=None) -> Path:
    """Write the document.

    values    {"BUILDING EXTERIOR:": {"Foundation": "..."} }, and for a
              repeating interior, {"Common Areas": {"Walls": "..."}}
    interiors the tenancy names, in the order they should appear
    """
    doc = Document(template_path)

    for name, copies in _repeat_interior(doc, interiors or []):
        for element in copies:
            _fill_placeholder(element, name)

    current = None
    for para in doc.paragraphs:
        if is_heading(para):
            title = _text(para).strip()
            current = None
            for key in (title, title.rstrip(":"), title.split("–")[-1].strip().rstrip(":")):
                if key in values:
                    current = values[key]
                    break
        elif is_field(para):
            wanted = (current or {}).get(field_label(para))
            _set_value(para, wanted if wanted else NOT_FOUND)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
