"""Build Mark's Photo.docx from a photo manifest, clone-and-fill from his template.

Ships (MANIFEST row). Pure Python: runs identically on Windows. Measured furniture
2026-08-09 from the delivered Mark Folder Template Photo.docx: 14 tables, each
3 rows x 2 columns (image cell, caption cell), image ~4.0in wide in the left
cell, centered caption in the right cell.

Body element order is NOT a simple [tbl][p][tbl] alternation. The measured order
is:

    p p tbl p p p tbl p p p tbl ... tbl p p tbl p p tbl p p p tbl ... tbl p sectPr

Two leading paragraphs (the "SUBJECT PHOTOGRAPHS" heading and a spacer), then
14 tables, each pair of consecutive tables separated by a run of 2 or 3
paragraphs (one of which carries an explicit <w:br w:type="page"/> to force
the next table onto a new page), and a single trailing paragraph after the
last table before the section properties. Grow/shrink logic below reads
whatever separator actually sits between two tables at clone time instead of
assuming a fixed paragraph count, so it works regardless of which run length
appears at that point in the template.
"""
import argparse
import copy
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image, ExifTags

PHOTOS_PER_TABLE = 3

# The box a photograph is fitted inside, in inches. Both numbers, not just the
# width.
#
# This is insurance, and it is not what caused the extra page. Corrected
# 2026-08-23 after Spenser checked: there is no portrait photograph anywhere in
# the photo documents he delivers. Mason City fifty images, 217 East 37th
# Street twenty-four, Burlington six, every one of them landscape. The extra
# page was the empty paragraph after a full last table, and it is dealt with in
# `_drop_trailing_blank_paragraphs`.
#
# What this still earns its place for: the shipped practice job does contain
# portrait photographs, and on the width alone one came out 4.00 x 5.33 in a
# row 2.93 inches tall. If a portrait photograph ever does reach a real job the
# document should be merely smaller, not broken.
#
# A landscape photograph is limited by the width and comes out 4.00 x 3.00,
# which is the size his delivered reports already use, so nothing that was
# right has moved.
IMAGE_WIDTH_IN = 4.0
IMAGE_MAX_HEIGHT_IN = 3.0

_DATETIME_TAG = next(k for k, v in ExifTags.TAGS.items() if v == "DateTimeOriginal")


def exif_order(paths):
    """Default ordering: EXIF capture time, falling back to filename."""
    def key(p: Path):
        try:
            exif = Image.open(p).getexif()
            stamp = exif.get(_DATETIME_TAG) or exif.get(306)  # 306 = DateTime
            if stamp:
                return (0, str(stamp), p.name.lower())
        except Exception:
            pass
        return (1, "", p.name.lower())
    return sorted(paths, key=key)


DEFAULT_OUTPUT_BASE = "Photo (RRF App)"


def next_output_name(photos_dir: Path, base: str = DEFAULT_OUTPUT_BASE) -> str:
    """Pick a save-as name for the built file that never collides with anything
    already in photos_dir, including Mark's own hand-built Photo.docx.

    The base name is the caller's now. It is `City_Address Photos (Complete)`
    for a job whose brief yields both values, and the old `Photo (RRF App)`
    remains the default so the command-line entry point and anything else that
    does not know a city still works. The counting is unchanged: if the name is
    taken it counts up until it finds one that is free, and it never returns an
    existing filename.
    """
    candidate = f"{base}.docx"
    n = 2
    while (photos_dir / candidate).exists():
        candidate = f"{base} {n}.docx"
        n += 1
    return candidate


def _fitted_size(image_path: Path):
    """The size to place this photograph at, fitted inside the box.

    Both edges are honoured and the shape is kept. A photograph wider than it
    is tall is limited by the width, exactly as before; a tall one is limited
    by the height instead and comes out narrower.

    The file measured here is the one being embedded. When the app supplies a
    preparer it has already applied the EXIF orientation, so a photograph taken
    on its side is measured the way Word will draw it rather than the way the
    camera happened to store it.

    A file PIL cannot open falls back to the width alone, which is what this
    did for everything until now. A photograph that cannot be measured is not a
    reason to fail a build: python-docx opens it a moment later and raises
    something the screen can show.
    """
    try:
        with Image.open(image_path) as opened:
            pixel_w, pixel_h = opened.size
    except Exception:
        return Inches(IMAGE_WIDTH_IN), None
    if pixel_w <= 0 or pixel_h <= 0:
        return Inches(IMAGE_WIDTH_IN), None

    shape = pixel_w / pixel_h
    width = IMAGE_WIDTH_IN
    height = width / shape
    if height > IMAGE_MAX_HEIGHT_IN:
        height = IMAGE_MAX_HEIGHT_IN
        width = height * shape
    return Inches(width), Inches(height)


def _fill_cell_image(cell, image_path: Path):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width, height = _fitted_size(image_path)
    p.add_run().add_picture(str(image_path), width=width, height=height)


def _fill_cell_caption(cell, caption: str):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    p.add_run(caption)


def _rewrite_copyright_year(para, year_str: str):
    """Replace the 4-digit year right after '©' in a paragraph, surgically.

    The footer paragraph mixes literal text runs with Word field-boundary
    runs (empty strings) for the "Page N of M" fields, and the copyright
    year itself can be split across run boundaries (e.g. '\\t©' + '202' +
    '2'). Rebuilding the paragraph text or clearing runs would destroy the
    field runs and break pagination, so this only overwrites the exact
    character positions that make up the matched year -- every other run,
    including empty field-boundary runs, is left byte-for-byte untouched.
    If no 4-digit year follows a '©' in this paragraph, nothing is changed.
    """
    runs = para.runs
    texts = [r.text for r in runs]
    full_text = "".join(texts)
    match = re.search(r"©(\d{4})", full_text)
    if not match:
        return
    start, end = match.start(1), match.end(1)

    offset = 0
    for run, text in zip(runs, texts):
        run_start, run_end = offset, offset + len(text)
        offset = run_end
        lo, hi = max(run_start, start), min(run_end, end)
        if lo < hi:
            local_lo, local_hi = lo - run_start, hi - run_start
            year_lo, year_hi = lo - start, hi - start
            run.text = text[:local_lo] + year_str[year_lo:year_hi] + text[local_hi:]


_HEADER_FOOTER_VARIANTS = (
    "header", "first_page_header", "even_page_header",
    "footer", "first_page_footer", "even_page_footer",
)


def _set_copyright_year(doc: Document, year: int):
    """Rewrite the copyright year everywhere it appears, in every section.

    python-docx's `section.footer` is only the *default* footer part. Word
    documents can carry up to three variants per section -- default, first
    page, and even page -- selected by `titlePg` / odd-even settings, and
    they are backed by *separate* XML parts (footer1/2/3.xml in this
    template), not shared text. The measured template uses this: its
    footer2.xml (default) said "(c)2022" and its footer3.xml (first page,
    which is what actually renders on page 1 because titlePg is set) said
    "(c)2023" -- two different stale years in two different parts. Only
    touching `section.footer` leaves the first-page one wrong.

    This iterates all three footer variants and, for completeness, all
    three header variants too (the template's headers only carry the
    "SUBJECT PHOTOGRAPHS- CONTINUED" running title with no year in them
    today, so this is a no-op there, but a future template revision that
    puts a year in a header is covered without further code changes).
    Sections that inherit a variant from a previous section
    (`is_linked_to_previous`) still resolve to the correct underlying part
    via python-docx, so processing every section is redundant-but-safe, not
    wrong -- the rewrite is idempotent.
    """
    year_str = str(int(year))
    if len(year_str) != 4:
        return  # only the 4-digit-year case measured in the template is handled
    for section in doc.sections:
        for variant in _HEADER_FOOTER_VARIANTS:
            part = getattr(section, variant)
            for para in part.paragraphs:
                _rewrite_copyright_year(para, year_str)


def _between(a, b):
    """Sibling elements strictly between elements a and b (both exclusive)."""
    out = []
    node = a.getnext()
    while node is not None and node is not b:
        out.append(node)
        node = node.getnext()
    return out


def _grow_tables(doc: Document, needed: int):
    """Append cloned [separator][table] units until len(doc.tables) == needed.

    The separator cloned each time is whatever run of paragraphs actually
    separates the current last two tables, so growth beyond the template's
    own 14 tables reuses the template's real page-break furniture rather
    than a guessed shape.
    """
    while len(doc.tables) < needed:
        tables_el = [t._element for t in doc.tables]
        last_tbl_el = tables_el[-1]
        prev_tbl_el = tables_el[-2] if len(tables_el) > 1 else None
        sep_unit = _between(prev_tbl_el, last_tbl_el) if prev_tbl_el is not None else []

        anchor = last_tbl_el
        for el in sep_unit:
            clone = copy.deepcopy(el)
            anchor.addnext(clone)
            anchor = clone
        new_tbl = copy.deepcopy(last_tbl_el)
        anchor.addnext(new_tbl)


def _embedded_sectPr(el):
    """The <w:sectPr> embedded in a paragraph's <w:pPr>, if any -- i.e. this
    paragraph is where a document section ends. That sectPr is where header
    and footer part references actually live; the template's very first
    inter-table separator (between table 1 and table 2) carries one, ending
    a short "title" section and starting the section that every later table
    lives in and inherits its footer from. Deleting that paragraph merges
    the two sections into one whose sectPr never declared a footer
    reference, silently erasing the footer for the entire document.
    """
    if not el.tag.endswith("}p"):
        return None
    ppr = el.find(qn("w:pPr"))
    if ppr is None:
        return None
    return ppr.find(qn("w:sectPr"))


def _find_embedded_sectPr(body):
    """The (paragraph, sectPr) pair for the template's embedded section
    break, if the paragraph carrying it is still present in `body`."""
    for el in body:
        sect = _embedded_sectPr(el)
        if sect is not None:
            return el, sect
    return None, None


def _table_follows(el):
    """True if a <w:tbl> appears anywhere after `el` among its siblings."""
    node = el.getnext()
    while node is not None:
        if node.tag.endswith("}tbl"):
            return True
        node = node.getnext()
    return False


def _merge_orphaned_section(doc: Document):
    """If shrinking removed every table that used to follow the template's
    embedded section break, eliminate the now-vestigial second section
    entirely rather than try to neutralize its page break.

    First attempt (round 2, take 1) was to leave the section-break
    paragraph in place and flip its type from the OOXML default "nextPage"
    to "continuous". Word rendering (checked by the coordinator, not
    available in this environment) showed that was NOT sufficient: Word
    still rendered an extra, blank "continuation" page carrying the
    second section's "SUBJECT PHOTOGRAPHS- CONTINUED" header, apparently
    because that section still legitimately existed (with its own
    titlePg-driven first-page header reference) even though it had a
    "continuous" break and zero real content.

    The robust fix is to remove the section boundary, not just soften it:
    promote the orphaned section's header/footer/page-setup references
    (the ones the *title* page actually uses -- first-page footer with the
    correct year, no CONTINUED header) onto the document's own final
    sectPr, replacing whatever the true final section declared, then
    delete the now-fully-redundant paragraph. The result is a single
    section using exactly the properties a genuinely one-table document
    should have, so there is no vestigial section left for Word to give a
    page to.

    A no-op if the section-break paragraph isn't present (never happens
    once `_shrink_tables` has run against this template) or if a table
    still follows it (the normal case for any kept-table count >= 2,
    where the second section still has real content and must be left
    exactly as the template defines it).
    """
    body = doc.element.body
    holder, sect = _find_embedded_sectPr(body)
    if holder is None or _table_follows(holder):
        return

    final_sect = body.find(qn("w:sectPr"))
    promoted = copy.deepcopy(sect)
    final_sect.addprevious(promoted)
    body.remove(final_sect)
    holder.getparent().remove(holder)


def _shrink_tables(doc: Document, needed: int):
    """Drop tables from the end, each together with the separator paragraphs
    immediately preceding it, until len(doc.tables) == needed.

    Removing the preceding separator along with each dropped table (rather
    than the trailing one) guarantees the new last kept table is followed
    directly by whatever originally followed the very last template table
    (a single trailing paragraph, then sectPr) -- never by a leftover
    page-break paragraph that would produce a trailing blank page.

    One exception: a separator paragraph that carries an embedded section
    break (see `_embedded_sectPr`) is never deleted here, because it is
    where this document's header/footer part references are actually
    declared -- blindly removing it merges sections and wipes the footer
    for every page, which is worse than the blank page this function
    otherwise avoids. If that protection leaves the paragraph guarding an
    empty, now-tableless section, `_merge_orphaned_section` cleans it up
    properly afterward (see its docstring for why a plain "make it
    continuous" fix was not enough).
    """
    while len(doc.tables) > needed:
        tables_el = [t._element for t in doc.tables]
        last_tbl_el = tables_el[-1]
        prev_tbl_el = tables_el[-2] if len(tables_el) > 1 else None
        if prev_tbl_el is not None:
            for el in _between(prev_tbl_el, last_tbl_el):
                if _embedded_sectPr(el) is not None:
                    continue
                el.getparent().remove(el)
        last_tbl_el.getparent().remove(last_tbl_el)
    _merge_orphaned_section(doc)


def _drop_trailing_blank_paragraphs(doc) -> int:
    """Remove the empty paragraphs sitting after the last table.

    This is the extra page. The template declares each row 2.926 inches tall
    and the app puts a 3.00 inch photograph in it, so every row grows past its
    own declared height and a full table of three ends up about a fifth of an
    inch too tall for the page. The empty paragraph that follows the last table
    is then one line too many, and it lands on a page of its own.

    Measured on 2026-08-23: twelve, nine and three photographs all end with a
    full table and all produced the extra page. Eleven did not, because its last
    row is empty and stays at the declared height. Mark's own delivered Mason
    City report ends on a partial table for the same reason and has no extra
    page, which is why this was never visible in the corpus.

    Removing the paragraph rather than shrinking the photographs is the smaller
    change: nothing about the page he prints moves, and 4.00 x 3.00 is the size
    his delivered reports already use. The `sectPr` stays where it is, so the
    document still ends the way Word expects.
    """
    body = doc.element.body
    removed = 0
    for node in list(body)[::-1]:
        tag = node.tag.split("}")[1]
        if tag == "sectPr":
            continue                     # the section always stays
        if tag == "p" and not "".join(node.itertext()).strip():
            body.remove(node)
            removed += 1
            continue
        break                            # stop at the first thing that prints
    return removed


def build_photo_docx(manifest_path: Path, template_path: Path,
                     prepare=None, out_base: str = DEFAULT_OUTPUT_BASE) -> Path:
    """Build the document. `prepare` decides what bytes actually go in.

    Without it the original file is embedded, which is what the command-line
    entry point still does. The app passes a preparer that hands back a
    temporary RGB JPEG: capped at 1,600 pixels, oriented by its EXIF, quality
    85. That is what keeps a sixty-photo report a sensible size, and it is also
    what lets a HEIC build at all, because python-docx cannot embed one.

    The originals are never given to `prepare` to modify. It returns a path to
    a copy somewhere else and the copies are deleted by the caller afterwards.
    """
    manifest = json.loads(Path(manifest_path).read_text())
    photos_dir = Path(manifest_path).parent
    # Photos cut from the report are skipped before anything is counted, so
    # the table count follows what is actually printed and no blank rows are
    # left behind. A missing "cut" means the photo is in, which is what every
    # manifest written before this feature says.
    photos = [e for e in manifest["photos"] if not e.get("cut")]
    doc = Document(str(template_path))

    needed = max(1, -(-len(photos) // PHOTOS_PER_TABLE))  # ceil
    _grow_tables(doc, needed)
    _shrink_tables(doc, needed)

    for i, entry in enumerate(photos):
        table = doc.tables[i // PHOTOS_PER_TABLE]
        row = table.rows[i % PHOTOS_PER_TABLE]
        source = photos_dir / entry["file"]
        _fill_cell_image(row.cells[0], prepare(source) if prepare else source)
        _fill_cell_caption(row.cells[1], entry.get("caption", ""))

    if manifest.get("report_year"):
        _set_copyright_year(doc, int(manifest["report_year"]))

    # Last, after every table is filled and trimmed, so what is trailing is
    # actually trailing.
    _drop_trailing_blank_paragraphs(doc)

    out = photos_dir / next_output_name(photos_dir, out_base)
    doc.save(str(out))
    return out


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--manifest", required=True, type=Path)
    ap.add_argument("--template", required=True, type=Path)
    args = ap.parse_args()
    out = build_photo_docx(args.manifest, args.template)
    print(out.name)


if __name__ == "__main__":
    main()
