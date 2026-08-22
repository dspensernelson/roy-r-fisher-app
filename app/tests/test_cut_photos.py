"""Cutting a photo from the report.

The whole feature is one optional key on one manifest entry. These tests are
mostly about what does NOT happen: the file on disk, the ordering, the
caption, and every other manifest value have to come through untouched.
"""
import copy
import json
import sys
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "server"))

import captions  # noqa: E402
import photos as photos_routes  # noqa: E402
from main import create_app  # noqa: E402

# Includes a key the app does not write itself, to prove nothing here
# rewrites a manifest it does not fully understand.
MANIFEST = {
    "job": "A job",
    "context": "1 Main Street, Davenport, Iowa · Retail",
    "report_year": 2026,
    "caption_style": "view",
    "ai_available": True,
    "photos": [
        {"file": "a.jpg", "caption": "View east"},
        {"file": "b.jpg", "caption": "The middle one"},
        {"file": "c.jpg", "caption": "Rear loading area"},
        {"file": "d.jpg", "caption": ""},
    ],
}


@pytest.fixture
def home(tmp_path, monkeypatch):
    monkeypatch.setenv("RRF_JOBS_HOME", str(tmp_path / "jobs"))
    jobs = tmp_path / "jobs"
    (jobs / "A job" / "Photos").mkdir(parents=True)
    for entry in MANIFEST["photos"]:
        (jobs / "A job" / "Photos" / entry["file"]).write_bytes(
            b"pretend jpeg for " + entry["file"].encode())
    (jobs / "A job" / "Photos" / "photo-manifest.json").write_text(
        json.dumps(MANIFEST, indent=2))
    return jobs


@pytest.fixture
def client(home):
    return TestClient(create_app())


def manifest_on_disk(home: Path) -> dict:
    return json.loads((home / "A job" / "Photos" / "photo-manifest.json").read_text())


def photo_bytes(home: Path) -> dict:
    d = home / "A job" / "Photos"
    return {p.name: p.read_bytes() for p in sorted(d.iterdir()) if p.suffix == ".jpg"}


# ------------------------------------------------------ the original file ---
def test_the_photo_on_disk_never_changes(client, home):
    before = photo_bytes(home)
    client.post("/api/jobs/A job/photos/b.jpg/cut")
    assert photo_bytes(home) == before
    client.post("/api/jobs/A job/photos/b.jpg/uncut")
    assert photo_bytes(home) == before


def test_no_file_is_moved_renamed_or_removed(client, home):
    d = home / "A job" / "Photos"
    before = sorted(p.name for p in d.iterdir())
    client.post("/api/jobs/A job/photos/a.jpg/cut")
    client.post("/api/jobs/A job/photos/c.jpg/cut")
    assert sorted(p.name for p in d.iterdir()) == before


# ------------------------------------------------------------- the state ---
def test_cutting_writes_only_the_cut_key(client, home):
    before = manifest_on_disk(home)
    client.post("/api/jobs/A job/photos/b.jpg/cut")
    after = manifest_on_disk(home)

    def masked(m):
        m = copy.deepcopy(m)
        for e in m["photos"]:
            e.pop("cut", None)
        return m

    assert masked(after) == masked(before)
    assert after["photos"][1]["cut"] is True
    assert all("cut" not in e for i, e in enumerate(after["photos"]) if i != 1)


def test_bringing_back_removes_the_key_rather_than_saving_false(client, home):
    client.post("/api/jobs/A job/photos/b.jpg/cut")
    client.post("/api/jobs/A job/photos/b.jpg/uncut")
    entry = manifest_on_disk(home)["photos"][1]
    assert "cut" not in entry, "a manifest should never carry a field meaning the default"


def test_unknown_manifest_fields_survive(client, home):
    client.post("/api/jobs/A job/photos/a.jpg/cut")
    after = manifest_on_disk(home)
    assert after["ai_available"] is True
    assert after["caption_style"] == "view"
    assert after["report_year"] == 2026
    assert after["context"] == MANIFEST["context"]
    assert after["job"] == "A job"


def test_the_caption_survives_being_cut_and_brought_back(client, home):
    client.post("/api/jobs/A job/photos/b.jpg/cut")
    assert manifest_on_disk(home)["photos"][1]["caption"] == "The middle one"
    client.post("/api/jobs/A job/photos/b.jpg/uncut")
    assert manifest_on_disk(home)["photos"][1]["caption"] == "The middle one"


# ---------------------------------------------------------- the ordering ---
def test_a_middle_photo_keeps_its_position(client, home):
    client.post("/api/jobs/A job/photos/b.jpg/cut")
    names = [e["file"] for e in manifest_on_disk(home)["photos"]]
    assert names == ["a.jpg", "b.jpg", "c.jpg", "d.jpg"]


def test_reordering_the_others_does_not_move_or_drop_the_cut_one(client, home):
    """The screen shuffles included photos between the slots they already
    hold, so the cut entry stays at index 1 whatever happens around it."""
    client.post("/api/jobs/A job/photos/b.jpg/cut")

    m = manifest_on_disk(home)
    slots = [i for i, e in enumerate(m["photos"]) if not e.get("cut")]
    items = [m["photos"][i] for i in slots]
    items.append(items.pop(0))                      # a.jpg to the end of the included
    for slot, item in zip(slots, items):
        m["photos"][slot] = item
    assert client.put("/api/jobs/A job/manifest", json=m).status_code == 200

    after = [e["file"] for e in manifest_on_disk(home)["photos"]]
    assert after == ["c.jpg", "b.jpg", "d.jpg", "a.jpg"]
    assert manifest_on_disk(home)["photos"][1]["file"] == "b.jpg"
    assert manifest_on_disk(home)["photos"][1]["cut"] is True

    client.post("/api/jobs/A job/photos/b.jpg/uncut")
    back = manifest_on_disk(home)["photos"]
    assert [e["file"] for e in back] == ["c.jpg", "b.jpg", "d.jpg", "a.jpg"]
    assert "cut" not in back[1]


# ------------------------------------------------------------- captions ----
def test_a_cut_photo_is_not_sent_for_captioning(client, home, monkeypatch):
    client.post("/api/jobs/A job/photos/d.jpg/cut")     # the only blank one
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")

    seen = []

    def fake(context, paths, style=None):
        seen.extend(p.name for p in paths)
        return {p.name: "written" for p in paths}, {"input": 100, "output": 20}

    monkeypatch.setattr(captions, "draft_captions", fake)
    client.post("/api/jobs/A job/captions")
    assert seen == [], "the only blank caption was on a cut photo"
    assert manifest_on_disk(home)["photos"][3]["caption"] == ""


def test_captions_are_written_for_the_included_blanks_only(client, home, monkeypatch):
    m = manifest_on_disk(home)
    for e in m["photos"]:
        e["caption"] = ""
    (home / "A job" / "Photos" / "photo-manifest.json").write_text(json.dumps(m, indent=2))
    client.post("/api/jobs/A job/photos/b.jpg/cut")

    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")
    seen = []

    def fake(context, paths, style=None):
        seen.extend(p.name for p in paths)
        return {p.name: "written" for p in paths}, {"input": 100, "output": 20}

    monkeypatch.setattr(captions, "draft_captions", fake)
    client.post("/api/jobs/A job/captions")

    assert sorted(seen) == ["a.jpg", "c.jpg", "d.jpg"]
    assert manifest_on_disk(home)["photos"][1]["caption"] == ""


def test_the_model_is_not_called_when_everything_is_cut(client, home, monkeypatch):
    for name in ["a.jpg", "b.jpg", "c.jpg", "d.jpg"]:
        client.post("/api/jobs/A job/photos/%s/cut" % name)
    monkeypatch.setenv("ANTHROPIC_API_KEY", "test-key")

    called = []
    monkeypatch.setattr(captions, "draft_captions",
                        lambda *a, **k: called.append(1) or {})

    client.post("/api/jobs/A job/captions")
    assert called == [], "nothing to caption means nothing is asked of the model"


# ---------------------------------------------------------------- build ----
def test_the_build_refuses_when_every_photo_is_cut(client, home):
    for name in ["a.jpg", "b.jpg", "c.jpg", "d.jpg"]:
        client.post("/api/jobs/A job/photos/%s/cut" % name)
    r = client.post("/api/jobs/A job/build")
    assert r.status_code == 400
    assert r.json()["detail"] == "Bring back at least one photo to build the report."


def test_the_engine_prints_only_included_photos_with_no_blank_rows(tmp_path):
    """Straight at the engine, because it reads the manifest off disk itself."""
    sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "engine"))
    from photo_pages import build_photo_docx, PHOTOS_PER_TABLE

    from conftest import TEMPLATE_DOCX
    if not TEMPLATE_DOCX.is_file():
        pytest.skip("template not on this machine")

    from PIL import Image
    for name in ["a.jpg", "b.jpg", "c.jpg", "d.jpg"]:
        Image.new("RGB", (60, 40), "grey").save(tmp_path / name)
    manifest = copy.deepcopy(MANIFEST)
    manifest["photos"][1]["cut"] = True             # b.jpg out
    mpath = tmp_path / "photo-manifest.json"
    mpath.write_text(json.dumps(manifest, indent=2))

    out = build_photo_docx(mpath, TEMPLATE_DOCX)
    from docx import Document
    doc = Document(str(out))

    rows = [r for t in doc.tables for r in t.rows]
    filled = [r for r in rows if r.cells[1].text.strip() or
              r.cells[0].paragraphs[0].runs]
    assert len(filled) == 3, "three included photos, three filled rows"
    text = " ".join(r.cells[1].text for r in rows)
    assert "The middle one" not in text
    assert "View east" in text and "Rear loading area" in text
    assert len(doc.tables) == max(1, -(-3 // PHOTOS_PER_TABLE))


def test_a_manifest_with_no_cut_field_builds_exactly_as_before(tmp_path):
    sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "engine"))
    from photo_pages import build_photo_docx
    from conftest import TEMPLATE_DOCX
    if not TEMPLATE_DOCX.is_file():
        pytest.skip("template not on this machine")

    from PIL import Image
    for name in ["a.jpg", "b.jpg", "c.jpg", "d.jpg"]:
        Image.new("RGB", (60, 40), "grey").save(tmp_path / name)
    mpath = tmp_path / "photo-manifest.json"
    mpath.write_text(json.dumps(MANIFEST, indent=2))     # no cut anywhere

    from docx import Document
    doc = Document(str(build_photo_docx(mpath, TEMPLATE_DOCX)))
    text = " ".join(r.cells[1].text for t in doc.tables for r in t.rows)
    for caption in ["View east", "The middle one", "Rear loading area"]:
        assert caption in text


# --------------------------------------------------------------- uploads ---
def test_a_newly_uploaded_photo_starts_included(client, home):
    client.post("/api/jobs/A job/photos/a.jpg/cut")
    from PIL import Image
    import io
    buf = io.BytesIO()
    Image.new("RGB", (40, 30), "grey").save(buf, format="JPEG")
    client.post("/api/jobs/A job/photos",
                files=[("files", ("new.jpg", buf.getvalue(), "image/jpeg"))])
    entry = [e for e in manifest_on_disk(home)["photos"] if e["file"] == "new.jpg"][0]
    assert "cut" not in entry


def test_a_photo_dropped_into_the_folder_starts_included(client, home):
    client.post("/api/jobs/A job/photos/a.jpg/cut")
    (home / "A job" / "Photos" / "dropped.jpg").write_bytes(b"straight from the camera")
    body = client.get("/api/jobs/A job/manifest").json()
    entry = [e for e in body["photos"] if e["file"] == "dropped.jpg"][0]
    assert "cut" not in entry


# ------------------------------------------------------------- refusals ----
def test_a_photo_that_is_not_in_the_job_is_refused(client, home):
    assert client.post("/api/jobs/A job/photos/nope.jpg/cut").status_code == 404


def test_a_traversing_file_name_is_refused(client, home):
    r = client.post("/api/jobs/A job/photos/..%2F..%2Fevil.jpg/cut")
    assert r.status_code in (400, 404)


def test_a_manifest_carrying_a_bad_cut_value_is_refused(client, home):
    m = manifest_on_disk(home)
    m["photos"][0]["cut"] = "yes please"
    r = client.put("/api/jobs/A job/manifest", json=m)
    assert r.status_code == 400
    assert "true or false" in r.json()["detail"]


def test_included_is_the_one_rule_everything_uses():
    assert photos_routes.included({"photos": [
        {"file": "a.jpg", "caption": ""},
        {"file": "b.jpg", "caption": "", "cut": True},
        {"file": "c.jpg", "caption": "", "cut": False},
    ]}) == [{"file": "a.jpg", "caption": ""}, {"file": "c.jpg", "caption": "", "cut": False}]
