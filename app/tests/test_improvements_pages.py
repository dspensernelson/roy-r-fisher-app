"""The Description of Improvements document.

The template is the layout and the truth, so most of these read the template
rather than a list written here. A test that restates the layout would just be
a second copy of it, and a copy drifts.

One test reads Mark's own Blaul Lofts file. That is his client's work, it is
not in this repository, and the test skips and says so on a machine without it.
"""
import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from conftest import CORPUS  # noqa: E402
from engine import improvements_pages as imp  # noqa: E402

TEMPLATE = Path(__file__).resolve().parents[1] / "templates" / "Improvements.docx"
BLAUL = (CORPUS / "Description of Improvement Examples 8.17.26"
         / "Description of Improvements - Blaul Lofts.docx")

needs_blaul = pytest.mark.skipif(
    not BLAUL.is_file(),
    reason=f"Mark's delivered file is private and not on this machine: {BLAUL.name}")

EXTERIOR_SPINE = ["Foundation", "Exterior Walls", "Roof", "Windows"]


def _labels(path):
    doc = Document(path)
    return [imp.field_label(p) for p in doc.paragraphs if imp.is_field(p)]


def _headings(path):
    doc = Document(path)
    return [imp._text(p).strip() for p in doc.paragraphs if imp.is_heading(p)]


def test_the_template_ships_with_every_value_empty():
    doc = Document(TEMPLATE)
    for para in doc.paragraphs:
        if imp.is_field(para):
            after = imp._text(para).strip().rstrip(":")
            assert after == imp.field_label(para), f"{after!r} ships with a value"


def test_the_template_carries_the_measured_exterior_spine():
    """The one part of the layout that is identical across all three of Mark's
    apartment and mixed-use reports, in this order. Measured 2026-08-28."""
    blocks = dict(imp.read_shape(TEMPLATE))
    assert blocks["BUILDING EXTERIOR:"] == EXTERIOR_SPINE


def test_the_interior_block_repeats_once_per_tenancy(tmp_path):
    names = ["Common Areas", "Commercial Suite", "Apartment Units"]
    out = imp.build_improvements_docx({}, TEMPLATE, tmp_path / "out.docx",
                                      interiors=names)
    headings = _headings(out)
    for name in names:
        assert f"BUILDING INTERIOR – {name}:" in headings
    assert not any(imp.PLACEHOLDER in h for h in headings)


@pytest.mark.parametrize("count", [1, 3, 6])
def test_any_number_of_tenancies_builds(tmp_path, count):
    names = [f"Part {n}" for n in range(count)]
    out = imp.build_improvements_docx({}, TEMPLATE, tmp_path / f"{count}.docx",
                                      interiors=names)
    headings = _headings(out)
    assert sum(h.startswith("BUILDING INTERIOR") for h in headings) == count


def test_no_tenancies_removes_the_repeating_block(tmp_path):
    out = imp.build_improvements_docx({}, TEMPLATE, tmp_path / "none.docx",
                                      interiors=[])
    assert not any(h.startswith("BUILDING INTERIOR") for h in _headings(out))


def test_a_value_with_no_source_says_so_rather_than_going_blank(tmp_path):
    """A blank costs Mark ten seconds. A silently missing field costs him a
    wrong report, because nothing on the page tells him to look."""
    out = imp.build_improvements_docx({}, TEMPLATE, tmp_path / "blank.docx")
    doc = Document(out)
    fields = [p for p in doc.paragraphs if imp.is_field(p)]
    assert fields
    for para in fields:
        assert imp.NOT_FOUND in imp._text(para)


def test_a_value_lands_against_its_own_label(tmp_path):
    out = imp.build_improvements_docx(
        {"BUILDING EXTERIOR": {"Roof": "Flat, rubber membrane."}},
        TEMPLATE, tmp_path / "one.docx")
    doc = Document(out)
    got = {imp.field_label(p): imp._text(p) for p in doc.paragraphs if imp.is_field(p)}
    assert "Flat, rubber membrane." in got["Roof"]
    assert imp.NOT_FOUND in got["Windows"]


def test_the_page_geometry_survives_being_filled(tmp_path):
    """A template can keep its margins inside a trailing empty paragraph.
    Deleting paragraphs can take the page size with it and nothing says so."""
    before = Document(TEMPLATE).sections
    out = imp.build_improvements_docx({}, TEMPLATE, tmp_path / "geom.docx",
                                      interiors=["A", "B"])
    after = Document(out).sections
    assert len(after) == len(before)
    assert after[0].page_width == before[0].page_width
    assert after[0].left_margin == before[0].left_margin


def test_the_title_repeats_by_itself_rather_than_being_typed(tmp_path):
    """Mark types his page title on each page, so it lands in the wrong place
    the moment the text length changes. Ours is a real running header."""
    out = imp.build_improvements_docx({}, TEMPLATE, tmp_path / "hdr.docx",
                                      interiors=["A"])
    doc = Document(out)
    header = " ".join(p.text for p in doc.sections[0].header.paragraphs)
    assert "DESCRIPTION OF IMPROVEMENTS" in header
    body = [p.text.strip() for p in doc.paragraphs]
    assert "DESCRIPTION OF IMPROVEMENTS" not in body


def test_a_heading_cannot_strand_at_the_foot_of_a_page():
    """Found by building a mockup: without this a heading ends one page and its
    own fields start the next. Mark's file cannot show this because his
    headings sit wherever his text happened to break."""
    doc = Document(TEMPLATE)
    for para in doc.paragraphs:
        if imp.is_heading(para):
            assert para.paragraph_format.keep_with_next, f"{para.text!r} can strand"


@needs_blaul
def test_our_exterior_labels_match_marks_own_file():
    blaul = _labels(BLAUL)
    ours = dict(imp.read_shape(TEMPLATE))["BUILDING EXTERIOR:"]
    positions = [blaul.index(label) for label in ours]
    assert positions == sorted(positions), "our order disagrees with Mark's"


@needs_blaul
def test_we_carry_the_two_blocks_blaul_leaves_out():
    """Approved 2026-08-28. Blaul omits both, and his own transcript carries
    the mechanical detail, so the omission is his rather than the source's."""
    blaul = _headings(BLAUL)
    ours = _headings(TEMPLATE)
    for block in ("MECHANICAL EQUIPMENT:", "SITE IMPROVEMENTS:"):
        assert block in ours
        assert block not in blaul
