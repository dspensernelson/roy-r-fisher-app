"""The document ends on its last photograph, not on a page of nothing.

Spenser found it: an extra page at the back of every built document, and the
page was empty.

The cause, measured on 2026-08-23. The template declares each table row 2.926
inches tall and the app places a 3.00 inch photograph in it, so every row grows
past its own declared height and a full table of three ends up roughly a fifth
of an inch too tall for the page. The empty paragraph that follows the last
table is then one line too many and lands on a page of its own.

The evidence that fixed the shape of it: twelve, nine and three photographs all
end on a full table and all showed the extra page. Eleven did not, because its
last row is empty and stays at the declared height. Mark's own delivered Mason
City report ends on a partial table for the same reason, which is why fifty
delivered photographs never showed this.

What was ruled out on the way, and is worth writing down so it is not chased
again: it is not padded rows. Twelve, nine and three carry exactly one row per
photograph. It is not portrait photographs; there is not one in any document he
delivers. It is not a trailing page break; ours has none.
"""
import json
import sys
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "engine"))
from photo_pages import PHOTOS_PER_TABLE, build_photo_docx  # noqa: E402

from conftest import TEMPLATE_DOCX, has_template


def a_job(tmp_path: Path, count: int) -> Path:
    """`count` landscape photographs, the only kind his reports contain."""
    photos = tmp_path / "JOB" / "Photos"
    photos.mkdir(parents=True)
    names = []
    for i in range(count):
        name = "photo-%02d.jpg" % (i + 1)
        Image.new("RGB", (1000, 750), (60, 110, 160)).save(photos / name)
        names.append(name)
    manifest = photos / "photo-manifest.json"
    manifest.write_text(json.dumps({
        "job": "JOB", "context": "1 Test Street, Davenport, Iowa",
        "report_year": 2026,
        "photos": [{"file": n, "caption": "View of the subject"} for n in names],
    }))
    return manifest


def tail_after_last_table(path: Path):
    """Everything in the body after the final table, by tag."""
    doc = Document(str(path))
    body = list(doc.element.body)
    last = body.index(doc.tables[-1]._tbl)
    return [k.tag.split("}")[1] for k in body[last + 1:]]


# A full last page is the case that broke; a partial one is the case that did
# not. Both are covered so a future change cannot fix one by breaking the other.
FULL_LAST_PAGE = [3, 6, 9, 12, 60]
PARTIAL_LAST_PAGE = [1, 2, 4, 5, 11]


@has_template
@pytest.mark.parametrize("count", FULL_LAST_PAGE + PARTIAL_LAST_PAGE)
def test_nothing_empty_trails_the_last_photograph(count, tmp_path):
    out = build_photo_docx(a_job(tmp_path, count), TEMPLATE_DOCX)
    assert tail_after_last_table(out) == ["sectPr"], (
        "an empty paragraph after the last table is the extra page")


@has_template
@pytest.mark.parametrize("count", FULL_LAST_PAGE + PARTIAL_LAST_PAGE)
def test_the_document_still_ends_the_way_word_expects(count, tmp_path):
    """The section stays. Removing it would be a different bug."""
    out = build_photo_docx(a_job(tmp_path, count), TEMPLATE_DOCX)
    doc = Document(str(out))
    assert len(doc.sections) >= 1
    assert list(doc.element.body)[-1].tag.endswith("}sectPr")


@has_template
@pytest.mark.parametrize("count", FULL_LAST_PAGE + PARTIAL_LAST_PAGE)
def test_every_photograph_is_still_in_the_document(count, tmp_path):
    """The whole point of the file. Nothing was dropped along with the blanks."""
    out = build_photo_docx(a_job(tmp_path, count), TEMPLATE_DOCX)
    doc = Document(str(out))
    assert len(doc.inline_shapes) == count
    assert len(doc.tables) == -(-count // PHOTOS_PER_TABLE)


@has_template
def test_a_caption_is_never_mistaken_for_a_blank(tmp_path):
    """Only empty paragraphs go. A trailing paragraph holding words stays."""
    manifest = a_job(tmp_path, 3)
    out = build_photo_docx(manifest, TEMPLATE_DOCX)
    doc = Document(str(out))
    captions = [c.text for t in doc.tables for r in t.rows for c in r.cells if c.text.strip()]
    assert len(captions) == 3


@has_template
def test_the_page_breaks_between_pages_are_untouched(tmp_path):
    """Only the tail was in scope. The breaks that separate pages stay."""
    out = build_photo_docx(a_job(tmp_path, 12), TEMPLATE_DOCX)
    doc = Document(str(out))
    breaks = sum(1 for k in doc.element.body
                 if k.tag.endswith("}p") and 'w:type="page"' in k.xml)
    sections = sum(1 for k in doc.element.body
                   if k.tag.endswith("}p") and "sectPr" in k.xml)
    assert breaks + sections == len(doc.tables) - 1, (
        "four tables need three separators between them")
