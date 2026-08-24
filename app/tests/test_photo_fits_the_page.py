"""A photograph has to fit the cell it is put in, whichever way up it is.

Found by Spenser on 2026-08-23, in the built document: a blank page at the back.

The cause was not the blank page. Every photograph was placed at a fixed four
inch width with its height left to follow the aspect ratio, so a landscape
photograph came out 4.00 x 3.00 and fitted the row, and a portrait one came out
4.00 x 5.33 and did not. The row is 2.93 inches. Five of the twelve photographs
in the shipped practice job are portrait, and each one pushed its page past the
printable area. The spare page at the end was the overflow arriving.

The measurement that settled it, from Mark's own delivered Mason City report:
fifty photographs, every one of them landscape, heights 2.74 to 3.31 inches,
widths 3.81 to 4.43. His corpus therefore says nothing about how he would like
a portrait photograph handled, because there is not one in it. What it does say
is the size a photograph on his page has always been, and that is the box a
portrait one now has to fit inside too.

Fitting inside a box rather than matching a width leaves every landscape
photograph exactly where it was, which is what makes this safe: the only
documents that change are the ones that were broken.
"""
import json
import sys
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Emu
from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "engine"))
from photo_pages import (IMAGE_MAX_HEIGHT_IN, IMAGE_WIDTH_IN,  # noqa: E402
                         PHOTOS_PER_TABLE, build_photo_docx)

from conftest import TEMPLATE_DOCX, has_template


def a_job(tmp_path: Path, sizes) -> Path:
    """A job whose photographs are the shapes given, as (width, height) pixels."""
    photos = tmp_path / "JOB" / "Photos"
    photos.mkdir(parents=True)
    names = []
    for i, (w, h) in enumerate(sizes):
        name = "photo-%02d.jpg" % (i + 1)
        Image.new("RGB", (w, h), (40, 90, 140)).save(photos / name)
        names.append(name)
    manifest = photos / "photo-manifest.json"
    manifest.write_text(json.dumps({
        "job": "JOB", "context": "1 Test Street, Davenport, Iowa",
        "report_year": 2026,
        "photos": [{"file": n, "caption": "View of the test subject"} for n in names],
    }))
    return manifest


def placed(path: Path):
    """Every image in the built document, as (width, height) in inches."""
    d = Document(str(path))
    return [(Emu(int(s.width)).inches, Emu(int(s.height)).inches)
            for s in d.inline_shapes]


LANDSCAPE = (1000, 750)
PORTRAIT = (750, 1000)
SQUARE = (900, 900)
VERY_TALL = (600, 1800)


@has_template
def test_a_portrait_photograph_fits_the_row(tmp_path):
    """The defect, at its smallest. 4.00 x 5.33 in a 2.93 inch row."""
    out = build_photo_docx(a_job(tmp_path, [PORTRAIT]), TEMPLATE_DOCX)
    (w, h), = placed(out)
    assert h <= IMAGE_MAX_HEIGHT_IN + 0.01, "a portrait photo is taller than its row"
    assert w <= IMAGE_WIDTH_IN + 0.01


@has_template
@pytest.mark.parametrize("shape", [LANDSCAPE, PORTRAIT, SQUARE, VERY_TALL])
def test_every_shape_of_photograph_fits_the_box(tmp_path, shape):
    out = build_photo_docx(a_job(tmp_path, [shape]), TEMPLATE_DOCX)
    (w, h), = placed(out)
    assert w <= IMAGE_WIDTH_IN + 0.01
    assert h <= IMAGE_MAX_HEIGHT_IN + 0.01


@has_template
@pytest.mark.parametrize("shape", [LANDSCAPE, PORTRAIT, SQUARE, VERY_TALL])
def test_the_photograph_is_never_stretched(tmp_path, shape):
    """Fitted, not squashed. An appraisal photograph that has been distorted is
    worse than one that is small."""
    out = build_photo_docx(a_job(tmp_path, [shape]), TEMPLATE_DOCX)
    (w, h), = placed(out)
    assert abs((w / h) - (shape[0] / shape[1])) < 0.01


@has_template
def test_landscape_photographs_are_exactly_where_they_were(tmp_path):
    """The size Mark's delivered reports have always used, unchanged.

    This is what makes the fix safe to ship: a job of landscape photographs
    builds the same document it built before.
    """
    out = build_photo_docx(a_job(tmp_path, [LANDSCAPE]), TEMPLATE_DOCX)
    (w, h), = placed(out)
    assert abs(w - IMAGE_WIDTH_IN) < 0.01
    assert abs(h - 3.0) < 0.01


@has_template
def test_a_page_of_photographs_fits_the_printable_height(tmp_path):
    """Three to a page, and the page has to hold them."""
    out = build_photo_docx(a_job(tmp_path, [PORTRAIT] * PHOTOS_PER_TABLE), TEMPLATE_DOCX)
    doc = Document(str(out))
    section = doc.sections[0]
    usable = Emu(int(section.page_height) - int(section.top_margin)
                 - int(section.bottom_margin)).inches
    stacked = sum(h for _, h in placed(out))
    assert stacked <= usable, "three photographs do not fit on one page"


@has_template
def test_the_shipped_practice_job_does_not_overflow(tmp_path):
    """The job in the package, which is where Spenser saw it.

    Its twelve photographs are five portrait and seven landscape, which is the
    mix that produced the spare page.
    """
    sizes = [LANDSCAPE, LANDSCAPE, PORTRAIT, PORTRAIT, PORTRAIT, PORTRAIT,
             LANDSCAPE, PORTRAIT, LANDSCAPE, LANDSCAPE, LANDSCAPE, LANDSCAPE]
    out = build_photo_docx(a_job(tmp_path, sizes), TEMPLATE_DOCX)
    doc = Document(str(out))
    section = doc.sections[0]
    usable = Emu(int(section.page_height) - int(section.top_margin)
                 - int(section.bottom_margin)).inches

    every = placed(out)
    assert len(every) == len(sizes)
    for n in range(0, len(every), PHOTOS_PER_TABLE):
        page = every[n:n + PHOTOS_PER_TABLE]
        assert sum(h for _, h in page) <= usable, "page %d overflows" % (n // PHOTOS_PER_TABLE + 1)
