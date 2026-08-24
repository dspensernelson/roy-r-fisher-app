import json
import math
import sys
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "engine"))
from photo_pages import build_photo_docx, exif_order, next_output_name  # noqa: E402

from conftest import TEMPLATE_DOCX, has_template


def make_photos(d: Path, n: int) -> list[str]:
    names = []
    for i in range(n):
        p = d / f"IMG_{5100 + i}.jpg"
        Image.new("RGB", (400, 300), (i * 30 % 255, 80, 90)).save(p)
        names.append(p.name)
    return names


def write_manifest(photos_dir: Path, names: list[str]) -> Path:
    m = photos_dir / "photo-manifest.json"
    m.write_text(json.dumps({
        "job": "TESTJOB", "context": "123 Test St, Davenport, Iowa",
        "report_year": 2026,
        "photos": [{"file": n, "caption": f"View of test subject {i}"} for i, n in enumerate(names)],
    }))
    return m


@has_template
def test_build_fills_three_per_page_and_trims_tables(tmp_path):
    names = make_photos(tmp_path, 5)
    out = build_photo_docx(write_manifest(tmp_path, names), TEMPLATE_DOCX)
    d = Document(str(out))
    assert len(d.tables) == 2                      # ceil(5/3) pages
    assert len(d.inline_shapes) == 5               # one image per photo
    # Five photographs, five rows. The sixth slot used to be an empty row on
    # the last page; Spenser found it on the 61-photo job, where it was two.
    assert [len(t.rows) for t in d.tables] == [3, 2]
    caps = [r.cells[1].text.strip() for t in d.tables for r in t.rows]
    assert len(caps) == 5
    assert "View of test subject 0" in caps[0]
    assert all(c for c in caps), "every row holds a caption"


@has_template
def test_build_never_overwrites(tmp_path):
    names = make_photos(tmp_path, 1)
    m = write_manifest(tmp_path, names)
    first = build_photo_docx(m, TEMPLATE_DOCX)
    second = build_photo_docx(m, TEMPLATE_DOCX)
    assert first.exists() and second.exists() and first != second


def test_next_output_name_is_fresh(tmp_path):
    (tmp_path / "Photo.docx").write_bytes(b"x")
    name = next_output_name(tmp_path)
    assert name.endswith(".docx") and not (tmp_path / name).exists()


def test_exif_order_falls_back_to_name(tmp_path):
    names = make_photos(tmp_path, 3)          # generated JPEGs carry no EXIF
    ordered = exif_order([tmp_path / n for n in reversed(names)])
    assert [p.name for p in ordered] == names


def _footer_texts(doc):
    """All (default, first-page, even-page) footer texts across every section.

    python-docx's `section.footer` is only the *default* footer part -- the
    template stores separate XML parts for the first-page and even-page
    footers, selected at render time because `titlePg` is set. A fix that
    only checks `.footer` would have passed while the first-page footer
    (what page 1 actually shows) was still stale -- exactly how the round-1
    fix shipped with that gap.
    """
    return [
        "".join(run.text for p in getattr(section, variant).paragraphs for run in p.runs)
        for section in doc.sections
        for variant in ("footer", "first_page_footer", "even_page_footer")
    ]


def _first_page_header_text(doc, section_index):
    section = doc.sections[section_index]
    return "".join(run.text for p in section.first_page_header.paragraphs for run in p.runs)


@has_template
def test_build_sets_footer_year_and_preserves_page_fields(tmp_path):
    names = make_photos(tmp_path, 1)
    m = write_manifest(tmp_path, names)  # report_year: 2026; template's footer parts say 2022 / 2023
    out = build_photo_docx(m, TEMPLATE_DOCX)
    d = Document(str(out))
    footers = _footer_texts(d)
    combined = "".join(footers)
    assert "2026" in combined
    assert "2022" not in combined and "2023" not in combined
    assert "Page" in combined
    assert "of" in combined
    # Finding A regression: the *first-page* footer is a separate XML part
    # from the default footer and carries its own, differently-stale year
    # ("2023" vs. the default footer's "2022"). It is what page 1 actually
    # renders (titlePg is set), so it must be rewritten too, not just the
    # default footer.
    first_page_footer_text = "".join(
        run.text for p in d.sections[0].first_page_footer.paragraphs for run in p.runs
    )
    assert "2026" in first_page_footer_text


@has_template
@pytest.mark.parametrize(
    "n, expect_two_sections",
    [(1, False), (2, False), (3, False), (4, True), (7, True)],
)
def test_build_page_shape_has_no_orphaned_continuation_section(tmp_path, n, expect_two_sections):
    """Structural proxy for Finding B: rendered page count must equal
    ceil(n/3) with no trailing blank/continuation page, and the "SUBJECT
    PHOTOGRAPHS- CONTINUED" running header must still appear on pages 2+
    when there are pages 2+.

    This cannot prove Word's actual rendered page count -- only a real
    render can, and this repo has no Word automation available -- but it
    pins the section/header shape that produces the right behavior: when
    every photo fits on the title page (needed == 1 table), the template's
    normally-vestigial second section must be fully merged away rather than
    left empty for Word to paginate; whenever a real continuation page is
    needed (needed >= 2 tables), the template's original two-section shape,
    with the CONTINUED header intact on the continuation section's own
    first page, must be preserved unchanged.
    """
    names = make_photos(tmp_path, n)
    m = write_manifest(tmp_path, names)
    out = build_photo_docx(m, TEMPLATE_DOCX)
    d = Document(str(out))
    assert len(d.tables) == math.ceil(n / 3)

    if expect_two_sections:
        assert len(d.sections) == 2
        assert "CONTINUED" in _first_page_header_text(d, 1)
    else:
        assert len(d.sections) == 1
    # Page 1 is always section 0's own first page; it must never show the
    # continuation header, regardless of whether a second section exists.
    assert "CONTINUED" not in _first_page_header_text(d, 0)
