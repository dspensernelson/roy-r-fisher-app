import json
import shutil
import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "engine"))
from photo_pages import build_photo_docx, exif_order  # noqa: E402

from conftest import GOLDEN_DELIVERED, GOLDEN_PHOTOS, TEMPLATE_DOCX, has_golden_corpus

MASON_PHOTOS = GOLDEN_PHOTOS
DELIVERED = GOLDEN_DELIVERED


@has_golden_corpus
def test_mason_city_golden(tmp_path):
    raws = exif_order(sorted(MASON_PHOTOS.glob("*.jpeg")))[:12]
    assert len(raws) == 12, "expected at least 12 raw Mason City jpegs"
    for p in raws:
        shutil.copy2(p, tmp_path / p.name)
    manifest = tmp_path / "photo-manifest.json"
    manifest.write_text(json.dumps({
        "job": "MASON CITY", "context": "4151 4th St SW, Mason City, Iowa",
        "report_year": 2026,
        "photos": [{"file": p.name, "caption": f"View of subject, photo {i + 1}"} for i, p in enumerate(raws)],
    }))
    out = build_photo_docx(manifest, TEMPLATE_DOCX)
    built, delivered = Document(str(out)), Document(str(DELIVERED))
    assert len(built.tables) == 4                          # 12 photos / 3 per page
    assert len(built.inline_shapes) == 12
    # Parity with Mark's artifact: same table geometry, same margins
    assert (len(built.tables[0].rows), len(built.tables[0].columns)) == \
           (len(delivered.tables[0].rows), len(delivered.tables[0].columns))
    assert built.sections[0].left_margin == delivered.sections[0].left_margin
