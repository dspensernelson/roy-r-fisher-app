"""The template that ships inside the app is the one the appraiser's PC will use.

Every other build test points at the private corpus copy, which a clone may
not have. These tests point at app/templates/Photo.docx, which is always
present, so the path that actually runs on the appraiser's machine is never the one
path with no test on it. No skip marker on purpose.
"""
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Emu
from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "engine"))
from photo_pages import build_photo_docx  # noqa: E402

SHIPPED = Path(__file__).resolve().parents[1] / "templates" / "Photo.docx"


def test_shipped_template_matches_measured_furniture():
    d = Document(str(SHIPPED))
    s = d.sections[0]
    assert round(Emu(s.page_width).inches, 1) == 8.5
    assert round(Emu(s.left_margin).inches, 1) == 0.9
    assert len(d.tables) == 14
    for t in d.tables:
        assert len(t.rows) == 3 and len(t.columns) == 2
    assert d.paragraphs[0].text.strip() == "SUBJECT PHOTOGRAPHS"
    assert len(d.inline_shapes) == 0


def test_shipped_template_builds_photo_pages(tmp_path):
    names = []
    for i in range(5):
        p = tmp_path / f"IMG_{5100 + i}.jpg"
        Image.new("RGB", (400, 300), (i * 30 % 255, 80, 90)).save(p)
        names.append(p.name)
    manifest = tmp_path / "photo-manifest.json"
    manifest.write_text(json.dumps({
        "job": "TESTJOB", "context": "123 Test St, Davenport, Iowa",
        "report_year": 2026,
        "photos": [{"file": n, "caption": f"View of test subject {i}"}
                   for i, n in enumerate(names)],
    }))
    out = build_photo_docx(manifest, SHIPPED)
    d = Document(str(out))
    assert len(d.tables) == 2          # ceil(5/3) pages
    assert len(d.inline_shapes) == 5   # one image per photo
