from docx import Document
from docx.shared import Emu

from conftest import TEMPLATE_DOCX, has_template


@has_template
def test_template_matches_measured_furniture():
    d = Document(str(TEMPLATE_DOCX))
    s = d.sections[0]
    assert round(Emu(s.page_width).inches, 1) == 8.5
    assert round(Emu(s.left_margin).inches, 1) == 0.9
    tables = d.tables
    assert len(tables) == 14
    for t in tables:
        assert len(t.rows) == 3 and len(t.columns) == 2
    assert d.paragraphs[0].text.strip() == "SUBJECT PHOTOGRAPHS"
    assert d.inline_shapes == [] or len(d.inline_shapes) == 0
