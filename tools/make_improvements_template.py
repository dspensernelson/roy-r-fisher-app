"""Seed app/templates/Improvements.docx once.

The .docx is the layout and the truth, not this script. Open it in Word and
change a label and the app follows, because the engine reads the template's own
shape rather than a list written in code. This file exists so the first version
is reproducible and so its measurements are written down somewhere.

Measured from Mark's own Blaul Lofts file, 2026-08-28:
  Letter page, 0.88in left and right margins, Times New Roman 12pt, justified.
  A field is one paragraph: bold label, colon, tab, then the value.
  The value column sits at 1.5in with a 1.5in hanging indent, so a value that
  wraps lines up under itself instead of under the label.
  One empty paragraph between fields. Two between blocks.

Two things here are corrections to his file rather than copies of it:
  His page title is typed by hand on each page, so it drifts the moment the
  text length changes. Ours is a real running header.
  His block headings can strand at the foot of a page away from their own
  fields. Ours carry keep_with_next.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

TITLE = "DESCRIPTION OF IMPROVEMENTS"
VALUE_COL = Inches(1.5)

# The repeating block's label is substituted per tenancy, per building or per
# use. Blaul uses three: Common Areas, Commercial Suite, Apartment Units.
INTERIOR_PLACEHOLDER = "[NAME]"

BLOCKS = [
    ("GENERAL:", []),
    ("BUILDING EXTERIOR:", ["Foundation", "Exterior Walls", "Roof", "Windows"]),
    (f"BUILDING INTERIOR – {INTERIOR_PLACEHOLDER}:",
     ["Walls", "Ceilings", "Floors", "Kitchens", "Bathrooms"]),
    ("MECHANICAL EQUIPMENT:", ["HVAC", "Electrical Service", "Common Area"]),
    ("SITE IMPROVEMENTS:", ["Parking", "Trash Removal", "Plantings", "Sidewalks"]),
    ("CONCLUSION:", []),
]


def build(out_path: Path) -> Path:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(0)

    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    # 0.95 top rather than Mark's 0.5: his page title is body text, ours is a
    # real header and needs the room above the body.
    section.top_margin = Inches(0.95)
    section.bottom_margin = Inches(0.5)
    section.left_margin = section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.5)

    head = section.header.paragraphs[0]
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = head.add_run(TITLE)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    for index, (title, fields) in enumerate(BLOCKS):
        if index:
            doc.add_paragraph()
            doc.add_paragraph()

        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        heading.paragraph_format.keep_with_next = True
        heading.add_run(title).bold = True

        spacer = doc.add_paragraph()
        spacer.paragraph_format.keep_with_next = True

        for position, label in enumerate(fields):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt = para.paragraph_format
            fmt.left_indent = VALUE_COL
            fmt.first_line_indent = -VALUE_COL
            fmt.keep_together = True
            para.add_run(f"{label}:").bold = True
            para.add_run("\t")
            if position < len(fields) - 1:
                doc.add_paragraph()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    target = Path(__file__).resolve().parents[1] / "app" / "templates" / "Improvements.docx"
    if target.exists():
        raise SystemExit(f"{target} already exists. It is the truth now; edit it in Word.")
    print(build(target))
